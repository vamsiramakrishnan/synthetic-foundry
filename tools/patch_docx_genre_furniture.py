from pathlib import Path

path = Path("src/worldloom/render/docx.py")
text = path.read_text()
marker = "\n\ndef _contents(document, ir: ArtifactIR, g: StyleGenome) -> None:  # type: ignore[no-untyped-def]\n"
helper = '''


def _genre_front_matter(document, ir: ArtifactIR, g: StyleGenome) -> None:  # type: ignore[no-untyped-def]
    """Render native control rows from existing ecology metadata only."""
    if ir.metadata.get("realism_profile") != "ecology/v1":
        return
    genre = ir.metadata.get("artifact_genre")
    if genre not in {"decision_memo", "controlled_document", "incident_rca"}:
        return

    from datetime import datetime
    from docx.shared import Pt

    created = ir.metadata.get("worldloom_created")
    date = datetime.fromisoformat(created).strftime("%d %b %Y") if created else ""
    if genre == "decision_memo":
        rows = (("From", ir.metadata.get("author", "")), ("Date", date), ("Subject", ir.title))
    else:
        rows = (
            ("Status", ir.metadata.get("lifecycle", "").replace("_", " ").title()),
            ("Revision", ir.metadata.get("revision", "")),
            ("Date", date),
        )
    rows = tuple((label, value) for label, value in rows if value)
    if not rows:
        return

    table = document.add_table(rows=0, cols=2)
    _apply_table_borders(table, "horizontal", g.rule_weight)
    _apply_cell_padding(table, _cell_padding_pt(g))
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(_heading_pt(g, _TS_BODY))
                    run.font.color.rgb = _rgb(g.colour_roles["body_text"])
'''
if marker not in text:
    raise SystemExit("DOCX contents seam changed")
text = text.replace(marker, helper + marker, 1)
old = "    _cover(document, ir, g)\n    _contents(document, ir, g)\n"
new = "    _cover(document, ir, g)\n    _genre_front_matter(document, ir, g)\n    _contents(document, ir, g)\n"
if old not in text:
    raise SystemExit("DOCX render seam changed")
path.write_text(text.replace(old, new, 1))
