"""Step 5.3: Word documents.

The interesting test is not that a ``.docx`` opens. It is that the Word document
and the Markdown document of the *same artifact* say the same thing — same
headings, same prose, same figures — because both are projections of one resolved
IR. A renderer that quietly dropped a section or rounded a number differently
would give the corpus two answers to one question, and every other test here
would still pass.
"""

from __future__ import annotations

import io
import zipfile
from datetime import datetime

import docx as python_docx
import pytest

from worldloom import MonthEndClose, RetailWorld, World
from worldloom.models import Cell, Chart, ChartKind, Column, Row, Table
from worldloom.narrative import DeterministicProvider, references
from worldloom.render import docx as docx_renderer
from worldloom.render import ooxml, slug_for
from worldloom.render.values import format_value

PERIOD = "2026-03"


@pytest.fixture(scope="module")
def rendered() -> World:
    world = RetailWorld(seed=8128).build().run(
        MonthEndClose(period=PERIOD, include_operational_incident=True)
    )
    return world.narrate(DeterministicProvider()).render("docx", "markdown")


def _files(world: World, suffix: str) -> dict[str, bytes]:
    return {r.artifact_id: r.payload for r in world._rendered if r.path.endswith(suffix)}


def _document(payload: bytes):  # type: ignore[no-untyped-def]
    return python_docx.Document(io.BytesIO(payload))


def _text(payload: bytes) -> str:
    document = _document(payload)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# What gets rendered
# ---------------------------------------------------------------------------


def test_every_document_shaped_artifact_becomes_a_word_file(rendered: World) -> None:
    produced = {
        rendered.artifact_intents.by_id(
            next(ir.intent_id for ir in rendered.artifact_irs if ir.id == artifact_id)
        ).artifact_type
        for artifact_id in _files(rendered, ".docx")
    }
    assert produced == docx_renderer.HANDLES & {
        i.artifact_type for i in rendered.artifact_intents
    }
    assert produced, "the episode should plan at least one document"


def test_the_workbook_is_not_rendered_as_a_document(rendered: World) -> None:
    """A spreadsheet flattened into Word loses every formula that made it a source."""
    assert "finance_workbook" not in docx_renderer.HANDLES

    workbooks = [
        i.id for i in rendered.artifact_intents if i.artifact_type == "finance_workbook"
    ]
    assert workbooks
    assert not set(workbooks) & set(_files(rendered, ".docx"))


def test_a_confluence_page_stays_in_confluence(rendered: World) -> None:
    """Rendering it as Word would assert a filing that does not exist."""
    assert "confluence_page" not in docx_renderer.HANDLES


def test_one_artifact_has_one_basename_across_formats(rendered: World) -> None:
    """A reader who finds the Word file should be able to guess its Markdown twin."""
    word = {r.artifact_id: r.path for r in rendered._rendered if r.path.endswith(".docx")}
    markdown = {r.artifact_id: r.path for r in rendered._rendered if r.path.endswith(".md")}
    for artifact_id, path in word.items():
        assert markdown[artifact_id].removesuffix(".md") == path.removesuffix(".docx")


def test_slug_covers_every_planned_type(rendered: World) -> None:
    for intent in rendered.artifact_intents:
        assert slug_for(intent.artifact_type) == slug_for(intent.artifact_type)
        assert "_" not in slug_for(intent.artifact_type)


# ---------------------------------------------------------------------------
# The two formats must agree
# ---------------------------------------------------------------------------


def test_word_and_markdown_carry_the_same_prose(rendered: World) -> None:
    """The claim that matters. Both are projections of one IR; neither may drift."""
    facts = {fact.id: fact for fact in rendered.facts}
    word = _files(rendered, ".docx")
    checked = 0

    for ir in rendered.artifact_irs:
        if ir.id not in word:
            continue
        body = _text(word[ir.id])
        for section in ir.sections:
            assert section.heading in body, f"{ir.id}: missing section {section.heading!r}"
            if section.body:
                resolved = references.substitute(section.body, facts)
                for sentence in resolved.split(". "):
                    fragment = sentence.strip().rstrip(".")
                    if len(fragment) > 20:
                        assert fragment in body, f"{ir.id}: prose dropped — {fragment!r}"
                        checked += 1

    assert checked > 10, f"only compared {checked} fragments"


def test_table_values_match_the_markdown_rendering(rendered: World) -> None:
    """Same cell, same characters — one `format_value`, so neither can round alone."""
    word = _files(rendered, ".docx")
    checked = 0
    for ir in rendered.artifact_irs:
        if ir.id not in word:
            continue
        body = _text(word[ir.id])
        for section in ir.sections:
            if section.table is None:
                continue
            for row in section.table.rows:
                for column in section.table.columns:
                    cell = row.cells.get(column.key)
                    if cell is None or cell.value is None:
                        continue
                    text = format_value(cell.value, column.number_format)
                    assert text in body, f"{ir.id}: cell {row.key}/{column.key} = {text!r} missing"
                    checked += 1
    assert checked > 20, f"only checked {checked} cells"


def test_a_hidden_section_is_written_but_labelled(rendered: World) -> None:
    """Hidden means not part of the readable surface, not undocumented.

    Markdown includes these, so Word must too — one artifact whose appendix
    exists in one format and not the other is two artifacts.
    """
    word = _files(rendered, ".docx")
    ir = next(
        ir for ir in rendered.artifact_irs
        if ir.id in word and any(s.hidden for s in ir.sections)
    )
    body = _text(word[ir.id])
    assert "Not part of the readable surface" in body
    for section in ir.sections:
        if section.hidden:
            assert section.heading in body


def test_no_fact_reference_survives_into_the_document(rendered: World) -> None:
    """An unsubstituted `{{fact:...}}` in a finished document is a broken document."""
    for payload in _files(rendered, ".docx").values():
        body = _text(payload)
        assert "{{fact:" not in body
        assert "[missing " not in body


def test_prose_still_contains_no_bare_number(rendered: World) -> None:
    """The arithmetic rule holds after rendering: figures arrive by substitution."""
    facts = {fact.id: fact for fact in rendered.facts}
    for ir in rendered.artifact_irs:
        for section in ir.sections:
            if section.body:
                assert not references.bare_numbers(section.body), (
                    f"{ir.id}/{section.heading} restated a figure"
                )
                # And the substituted form does contain figures, or the
                # references were decorative.
                assert references.substitute(section.body, facts) != section.body


# ---------------------------------------------------------------------------
# Determinism
# ---------------------------------------------------------------------------


def test_no_clock_reaches_the_document(rendered: World) -> None:
    """Checked as an invariant, not by rendering twice.

    A comparison test passes by luck whenever two runs land in the same second,
    which is how the same defect went unnoticed in the XLSX renderer until CI
    happened to straddle a boundary.
    """
    for ir in rendered.artifact_irs:
        payload = _files(rendered, ".docx").get(ir.id)
        if payload is None:
            continue

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            stamps = {info.date_time for info in archive.infolist()}
        assert stamps == {ooxml.EPOCH}, f"{ir.id}: archive carries wall-clock entries {stamps}"

        properties = _document(payload).core_properties
        expected = datetime.fromisoformat(ir.metadata["worldloom_created"])
        assert properties.created == expected
        assert properties.modified == expected
        assert properties.created.year == 2026, "the template's own date leaked through"


def test_the_document_date_is_when_the_artifact_was_written(rendered: World) -> None:
    """It must match the manifest, or a file contradicts the corpus about itself."""
    word = _files(rendered, ".docx")
    for artifact in rendered.artifacts:
        if artifact.id not in word:
            continue
        properties = _document(word[artifact.id]).core_properties
        assert properties.created == artifact.created_at


def test_rendering_twice_is_byte_identical(rendered: World) -> None:
    ir = next(ir for ir in rendered.artifact_irs if ir.id in _files(rendered, ".docx"))
    facts = {fact.id: fact for fact in rendered.facts}
    assert docx_renderer.render(ir, facts) == docx_renderer.render(ir, facts)


def test_acronyms_survive_the_title(rendered: World) -> None:
    """"Cfo Variance Memo" is the tell that a document was generated."""
    titles = {ir.title for ir in rendered.artifact_irs}
    assert "CFO Variance Memo" in titles
    assert not any("Cfo" in title for title in titles)


# ---------------------------------------------------------------------------
# Document furniture
# ---------------------------------------------------------------------------


def _one(rendered: World, fragment: str):  # type: ignore[no-untyped-def]
    item = next(r for r in rendered._rendered if fragment in r.path and r.path.endswith(".docx"))
    return _document(item.payload)


def test_pages_are_a4_with_corporate_margins(rendered: World) -> None:
    """The worlds this renders are not American."""
    for payload in _files(rendered, ".docx").values():
        section = _document(payload).sections[0]
        assert round(section.page_width.mm) == 210
        assert round(section.page_height.mm) == 297
        assert round(section.top_margin.mm) == 22


def test_the_header_names_the_world_and_the_document(rendered: World) -> None:
    document = _one(rendered, "incident-rca")
    header = document.sections[0].header.paragraphs[0].text
    assert rendered.company.name in header
    assert "Incident RCA" in header


def test_the_footer_counts_pages_with_a_field_not_a_literal(rendered: World) -> None:
    """A written page count is wrong the moment a section is added."""
    footer = _one(rendered, "incident-rca").sections[0].footer.paragraphs[0]
    xml = footer._element.xml
    assert "PAGE" in xml and "NUMPAGES" in xml
    assert "fldChar" in xml, "page number should be a Word field"


def test_long_documents_get_a_contents_field(rendered: World) -> None:
    """A field, so Word builds it from the headings actually present."""
    from worldloom.render.docx import _TOC_THRESHOLD

    for ir in rendered.artifact_irs:
        payload = _files(rendered, ".docx").get(ir.id)
        if payload is None:
            continue
        visible = [s for s in ir.sections if not s.hidden]
        xml = _document(payload).element.xml
        assert (" TOC " in xml) == (len(visible) >= _TOC_THRESHOLD), ir.id


def test_table_headers_are_shaded_and_negatives_are_marked(rendered: World) -> None:
    """A financial table nobody can scan is a financial table nobody reads."""
    from worldloom.compiler.style import genome
    from worldloom.rng import Rng

    document = _one(rendered, "cfo-variance-memo")
    # The genome's fill, not the module constant. This used to assert
    # `docx._HEADER_FILL` and was right to until the look became a property of
    # the world rather than of the renderer. Asserting the constant now would
    # only pass for whichever seed happened to sample the house palette, and
    # would quietly stop testing anything for every other seed — so it derives
    # the expected fill exactly as the renderer does.
    expected = genome(Rng(rendered.seed).derive("style")).colour_roles["header_fill"]
    assert expected in document.element.xml
    # Parenthesised, so the sign survives a black-and-white printer; the colour
    # is the second signal, never the only one.
    assert any("(" in cell.text for table in document.tables for row in table.rows for cell in row.cells)


def test_the_memo_carries_a_divisional_table_and_a_figure(rendered: World) -> None:
    """A variance memo without a table makes its reader hold four numbers in their
    head while reading a paragraph about them."""
    ir = next(
        ir for ir in rendered.artifact_irs
        if rendered.artifact_intents.by_id(ir.intent_id).artifact_type == "cfo_variance_memo"
    )
    section = next(s for s in ir.sections if s.heading == "Divisional summary")
    assert section.table is not None and section.charts

    body = _text_of(_one(rendered, "cfo-variance-memo"))
    assert "Figure — Revenue against plan by division" in body

    charts = _chart_parts(_payload(rendered, "cfo-variance-memo"))
    assert len(charts) == 1, "one declared chart should become one chart part"
    assert "<c:barChart>" in charts[0], "ChartKind.BAR should draw as a bar chart, not a column one"


def _text_of(document) -> str:  # type: ignore[no-untyped-def]
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _payload(rendered: World, fragment: str) -> bytes:
    return next(r.payload for r in rendered._rendered if fragment in r.path and r.path.endswith(".docx"))


def _chart_parts(payload: bytes) -> list[str]:
    """Every chart part's XML text, decoded — there is one file per declared
    chart (`word/charts/chart1.xml`, `chart2.xml`, ...)."""
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        return [
            archive.read(name).decode("utf-8")
            for name in sorted(archive.namelist())
            if name.startswith("word/charts/chart") and name.endswith(".xml")
        ]


def test_the_figure_plots_the_same_values_the_table_shows(rendered: World) -> None:
    """The chart is drawn from the cells beside it, so it cannot disagree —
    checked in the chart part's own cache, since a native chart's plotted
    values live there rather than as document text `_text_of` can see."""
    ir = next(
        ir for ir in rendered.artifact_irs
        if rendered.artifact_intents.by_id(ir.intent_id).artifact_type == "cfo_variance_memo"
    )
    section = next(s for s in ir.sections if s.heading == "Divisional summary")
    chart = section.charts[0]
    measure = chart.series[0]

    chart_xml = _chart_parts(_payload(rendered, "cfo-variance-memo"))[0]
    checked = 0
    for row in section.table.rows:
        cell = row.cells.get(measure)
        if cell is None or not isinstance(cell.value, (int, float)):
            continue
        assert repr(float(cell.value)) in chart_xml, f"{row.key}: {cell.value!r} missing from the chart cache"
        checked += 1
    assert checked > 0, "the chart's own table should have plottable rows"


# ---------------------------------------------------------------------------
# Native charts — no first-class chart API in python-docx, so this is a
# hand-built `c:chartSpace` package part (see `docx.py::_chart_xml`'s own
# docstring). The tests above check the one chart the shipped corpus
# declares; these check the general machinery every `Chart` kind shares.
# ---------------------------------------------------------------------------


def _table_and_chart(kind: ChartKind, *, by_row: bool = False, rows: list[str] | None = None) -> tuple[Table, Chart]:
    table = Table(
        key="pnl",
        title="P&L",
        columns=[
            Column(key="budget", label="Budget", number_format="#,##0;(#,##0)"),
            Column(key="actual", label="Actual", number_format="#,##0;(#,##0)"),
        ],
        rows=[
            Row(key="food", label="Food", cells={"budget": Cell(value=100.0), "actual": Cell(value=110.0)}),
            Row(key="apparel", label="Apparel", cells={"budget": Cell(value=200.0), "actual": Cell(value=190.0)}),
            Row(key="group", label="Group", emphasis=True,
                cells={"budget": Cell(value=300.0), "actual": Cell(value=300.0)}),
        ],
    )
    chart = Chart(
        key="test", title="Test chart", kind=kind, table="pnl",
        series=["budget", "actual"], rows=rows if rows is not None else [], by_row=by_row,
        category_axis="Division", value_axis="AUD thousands", note="A note on the chart.",
    )
    return table, chart


#: A fixed "written at" stamp for the byte-identity helpers below. Any value
#: will do; what matters is that it is the *same* value on both renders, so
#: the comparison is about the renderer rather than about the clock.
_CREATED = "2026-04-01T07:00:00+00:00"


def _rendered_chart(kind: ChartKind, **kwargs):  # type: ignore[no-untyped-def]
    from itertools import count

    import docx as _docx

    from worldloom.compiler.style import genome as _genome
    from worldloom.rng import Rng as _Rng

    table, chart = _table_and_chart(kind, **kwargs)
    document = _docx.Document()
    g = _genome(_Rng(0).derive("style"), archetype="house")
    docx_renderer._figure(document, chart, table, g, count(1))
    buffer = io.BytesIO()
    document.save(buffer)
    # Through the normaliser, because that is what the renderer ships and
    # `document.save` alone is not it: python-docx stamps every zip entry with
    # the wall clock, so two saves a second apart differ in the DOS timestamp
    # of each part while the document is identical. A byte-identity assertion
    # over un-normalised output is a claim about how fast the test ran, and it
    # duly passed for months and then failed in CI on a slower runner, in the
    # template-derived parts (theme, fonts, settings) that happened to cross
    # the second boundary. `render` applies exactly this call.
    # And `normalise` alone is not it either, which is the same lesson one level
    # deeper. Without a *created* stamp it fixes archive entry dates and leaves
    # every XML timestamp exactly where it was — deliberately, pinned by
    # `test_ooxml.test_without_a_created_stamp_only_the_archive_entries_are_fixed`.
    # A chart embeds a data-source workbook that `xlsxwriter` stamps with
    # `datetime.now(timezone.utc)`, so an un-stamped normalise still leaves a
    # live clock nested two containers down. Measured on the PPTX sibling of
    # this helper: two renders 2.2s apart differ by one byte, inside
    # `ppt/embeddings/Microsoft_Excel_Sheet1.xlsx`, and the test passes or fails
    # on whether the two saves land in the same second — which is why it only
    # ever failed under a loaded machine. A determinism test has to pin every
    # clock, not most of them.
    from worldloom.render import ooxml

    return ooxml.normalise(buffer.getvalue(), created=_CREATED), table, chart


@pytest.mark.parametrize("kind", [ChartKind.COLUMN, ChartKind.BAR, ChartKind.LINE, ChartKind.PIE])
def test_every_chart_kind_draws_a_native_chart_part(kind: ChartKind) -> None:
    payload, _, _ = _rendered_chart(kind)
    charts = _chart_parts(payload)
    assert len(charts) == 1
    tag = {
        ChartKind.COLUMN: "<c:barChart>",
        ChartKind.BAR: "<c:barChart>",
        ChartKind.LINE: "<c:lineChart>",
        ChartKind.PIE: "<c:pieChart>",
    }[kind]
    assert tag in charts[0]
    # A round trip through python-docx itself, proving the package is not
    # merely well-formed XML but a document Word's own reader (which
    # `python-docx` implements against) accepts without error.
    python_docx.Document(io.BytesIO(payload))


def test_empty_rows_means_every_row_that_is_not_a_subtotal() -> None:
    """`Chart.rows`'s own docstring: a chart that included the subtotal would
    double every bar. This is the defect the old ASCII-bar `_figure` had —
    `rows=[]` fell through to *every* row, subtotal included."""
    payload, _table, _chart = _rendered_chart(ChartKind.COLUMN, rows=[])
    chart_xml = _chart_parts(payload)[0]
    assert "Food" in chart_xml and "Apparel" in chart_xml
    assert "Group" not in chart_xml, "the subtotal row should not be plotted"


def test_by_row_reads_rows_as_series_and_columns_as_categories() -> None:
    """`Chart.by_row`'s own docstring: read the wrong way round, this renders
    without complaint as one point per series instead of one line per row."""
    payload, _table, _chart = _rendered_chart(ChartKind.COLUMN, by_row=True, rows=["food", "apparel"])
    chart_xml = _chart_parts(payload)[0]
    # Series names are now the row labels, not the column labels.
    assert "<c:v>Food</c:v>" in chart_xml and "<c:v>Apparel</c:v>" in chart_xml
    # Categories are now the column labels.
    assert "<c:v>Budget</c:v>" in chart_xml and "<c:v>Actual</c:v>" in chart_xml
    assert chart_xml.count("<c:ser>") == 2, "one series per plotted row"


def test_a_percentage_column_is_scaled_the_same_way_the_workbook_scales_it() -> None:
    """A percentage fact is stored as e.g. 24.94; a chart plotting it should
    show `0.2494`, the same `xlsx.py::render` scales a percent cell to
    before Excel's own percent format multiplies it back out."""
    table = Table(
        key="margins", title="Margins",
        columns=[Column(key="gm_pct", label="GM%", number_format="0.00%")],
        rows=[
            Row(key="food", label="Food", cells={"gm_pct": Cell(value=24.94)}),
            Row(key="apparel", label="Apparel", cells={"gm_pct": Cell(value=31.5)}),
        ],
    )
    chart = Chart(key="margin", title="Margin", kind=ChartKind.COLUMN, table="margins", series=["gm_pct"])
    document = python_docx.Document()
    from itertools import count as _count

    from worldloom.compiler.style import genome as _genome
    from worldloom.rng import Rng as _Rng

    g = _genome(_Rng(0).derive("style"), archetype="house")
    docx_renderer._figure(document, chart, table, g, _count(1))
    buffer = io.BytesIO()
    document.save(buffer)
    chart_xml = _chart_parts(buffer.getvalue())[0]
    assert repr(0.2494) in chart_xml
    assert repr(24.94) not in chart_xml


def test_a_pie_chart_plots_only_its_first_series() -> None:
    """`ChartKind.PIE`'s own docstring: composition of a single total, one
    series only — a second series would silently double the wedges."""
    payload, _, _ = _rendered_chart(ChartKind.PIE)
    chart_xml = _chart_parts(payload)[0]
    assert chart_xml.count("<c:ser>") == 1


def test_a_chart_reading_a_different_table_than_its_section_is_skipped() -> None:
    """`Chart`'s own docstring: every value it plots is a cell already in the
    table beside it. A `chart.table` naming some other table is not a value
    to draw from a guess — it is dropped."""
    table, chart = _table_and_chart(ChartKind.COLUMN)
    mismatched = chart.model_copy(update={"table": "not_this_table"})
    from worldloom.models import ArtifactSection

    section = ArtifactSection(heading="Test", table=table, charts=[mismatched])
    document = python_docx.Document()
    from worldloom.compiler.style import genome as _genome
    from worldloom.rng import Rng as _Rng

    g = _genome(_Rng(0).derive("style"), archetype="house")
    docx_renderer._section(document, section, {}, g)
    buffer = io.BytesIO()
    document.save(buffer)
    assert _chart_parts(buffer.getvalue()) == []


def test_no_chart_carries_a_linked_workbook() -> None:
    """No `c:externalData`, and no embedded `.xlsx` part — see `_chart_xml`'s
    own docstring for why: the values are already cached in the chart XML,
    and a linked workbook would be a second OOXML package (and a second
    clock) nested inside this one for no reader-visible benefit."""
    payload, _, _ = _rendered_chart(ChartKind.COLUMN)
    chart_xml = _chart_parts(payload)[0]
    assert "externalData" not in chart_xml
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        assert not any(name.endswith(".xlsx") for name in archive.namelist())


def test_two_charts_in_one_document_get_distinct_drawing_ids() -> None:
    """`wp:docPr`'s `id` must be unique across the whole document — a second
    chart reusing "1" is exactly the bug a per-section counter would produce."""
    from itertools import count

    table, chart_a = _table_and_chart(ChartKind.COLUMN)
    chart_b = chart_a.model_copy(update={"key": "second", "title": "Second chart"})
    document = python_docx.Document()
    from worldloom.compiler.style import genome as _genome
    from worldloom.rng import Rng as _Rng

    g = _genome(_Rng(0).derive("style"), archetype="house")
    chart_index = count(1)
    docx_renderer._figure(document, chart_a, table, g, chart_index)
    docx_renderer._figure(document, chart_b, table, g, chart_index)
    buffer = io.BytesIO()
    document.save(buffer)

    docs = python_docx.Document(io.BytesIO(buffer.getvalue()))
    ids = [
        el.get("id")
        for el in docs.element.body.iter()
        if el.tag.endswith("}docPr")
    ]
    assert len(ids) == 2 and len(set(ids)) == 2


def test_a_chart_document_renders_twice_byte_identical() -> None:
    payload_a, _table, _chart = _rendered_chart(ChartKind.LINE, by_row=True, rows=["food", "apparel"])
    payload_b, _, _ = _rendered_chart(ChartKind.LINE, by_row=True, rows=["food", "apparel"])
    assert payload_a == payload_b
