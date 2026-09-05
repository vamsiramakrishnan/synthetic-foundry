from io import BytesIO

from docx import Document

from worldloom.models import ArtifactIR
from worldloom.render.docx import render


def _document(genre: str) -> Document:
    ir = ArtifactIR(
        id="doc",
        intent_id="intent",
        title="Forecast decision",
        metadata={
            "realism_profile": "ecology/v1",
            "artifact_genre": genre,
            "author": "Asha Rao",
            "worldloom_created": "2026-03-31T17:00:00+00:00",
            "lifecycle": "reviewed",
            "revision": "3",
        },
    )
    return Document(BytesIO(render(ir)))


def _rows(document: Document) -> dict[str, str]:
    if not document.tables:
        return {}
    return {
        row.cells[0].text: row.cells[1].text
        for row in document.tables[0].rows
        if len(row.cells) >= 2
    }


def test_decision_memo_uses_existing_metadata_as_front_matter() -> None:
    rows = _rows(_document("decision_memo"))
    assert rows == {
        "From": "Asha Rao",
        "Date": "31 Mar 2026",
        "Subject": "Forecast decision",
    }


def test_controlled_document_exposes_status_revision_and_date() -> None:
    rows = _rows(_document("controlled_document"))
    assert rows == {
        "Status": "Reviewed",
        "Revision": "3",
        "Date": "31 Mar 2026",
    }


def test_incident_rca_uses_the_same_small_control_block() -> None:
    rows = _rows(_document("incident_rca"))
    assert rows == {
        "Status": "Reviewed",
        "Revision": "3",
        "Date": "31 Mar 2026",
    }


def test_legacy_document_has_no_genre_front_matter() -> None:
    ir = ArtifactIR(id="doc", intent_id="intent", title="Legacy")
    document = Document(BytesIO(render(ir)))
    assert document.tables == []
